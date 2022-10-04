import errno
import datetime
import json
import os
import inspect
from tempfile import tempdir
import tempfile
import win32com.client
from docx import Document
from PyPDF3 import PdfFileReader, PdfFileMerger

def delete_files(files):
    for file in files:
        if os.path.exists(file):
            os.remove(file)

def escapeLatexSpecialChars(text):
    return (text.replace('_', r'\_').replace('%', r'\%').replace('#', r'\#').replace('&', r'\&')
            .replace('$', r'\$')
            .replace('{', r'\{')
            .replace('}', r'\}')
            .replace('~', r'\textasciitilde{}')
            .replace('^', r'\textasciicircum{}')
            .replace('\\', r'\textbackslash{}')
            .replace('<', r'\textless{}')
            .replace('>', r'\textgreater{}'))


def createTOC(template, files, title=''):

    if files == []:
        return

    document = Document(template)
    document.add_page_break()

    i = 1

    for file in files:
        if file == 'template.docx' or file == 'z_Combined.pdf':
            continue

        if title != '':
            for paragraph in document.paragraphs:
                if "{TITLE}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{TITLE}", title)

        name = file if os.path.isdir(file) else os.path.splitext(file)[0]
        index = file.split()[0]

        document.add_heading(
            index + '.	' + ' '.join(file.split()[2:]), level=1)

        if i < len(files):
            document.add_page_break()
        i += 1

    return document


def updateTOC(file_name):
    script_dir = os.path.dirname(os.path.abspath(
        inspect.getfile(inspect.currentframe())))
    file_path = os.path.join(script_dir, file_name)
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(file_path)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()


# merge pdf
def mergePDFs(inputA, inputB,  pos):
    output = os.path.join(os.getcwd(), 'z_Combined.pdf')
    tmpFile = os.path.join(tempfile.gettempdir(), '.lyntouch-toc/tmp.pdf')
    merger = PdfFileMerger()
    merger.append(inputA)
    merger.merge(fileobj=inputB, position=pos)
    try:
        merger.write(tmpFile)
    except OSError as e:
        if e.errno != errno.ENOENT:
            raise

    if os.path.exists(output):
        os.remove(output)
    os.rename(tmpFile, output)
    merger.close()


def readJson(file_name):
    if os.path.isfile(file_name):
        with open(file_name, "r") as f:
            config = json.load(f)
            return config
    else:
        config = {
            "company": "WICL",
            "from":  "Weston International Capital Limited",
            "to": "Compnay",
            "subject": "Subject",
            "date": datetime.datetime.now().strftime("%d %B %Y")
        }
        # set config date to be today

        with open(file_name, "w") as f:
            json.dump(config, f, indent=4)
        return config
